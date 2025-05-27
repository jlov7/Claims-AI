import logging
import os
import uuid
import re
from typing import Optional, Dict, Any, List, Union
from fastapi import HTTPException
from docx import Document as DocxDocument
from docx.shared import Pt
from pathlib import Path

from core.config import Settings, get_settings
from models import DraftStrategyNoteRequest, QAPair
from langchain_community.chat_models import ChatOllama
from langchain_core.prompts import ChatPromptTemplate
from langchain_core.output_parsers import StrOutputParser
from langchain_core.runnables import RunnablePassthrough

# For tool calling - specific to draft agent
# from backend.agents.draft import draft_agent_node
# from backend.core.state import DraftAgentState
from langgraph.graph import StateGraph, END

# Import DocumentLoaderService and its getter

logger = logging.getLogger(__name__)

# Default output directory within the container/filesystem for generated drafts
DEFAULT_DRAFT_OUTPUT_DIR = Path("data/outputs/strategy_notes")

MAX_FILENAME_LEN = 100  # Max length for a filename stem


def create_draft_graph():
    """Creates and compiles the LangGraph for the draft agent."""
    workflow = StateGraph(DraftAgentState)
    workflow.add_node("draft_strategist", draft_agent_node)  # Using the imported node
    workflow.set_entry_point("draft_strategist")
    # For now, a simple graph that ends after the draft_strategist node.
    # It could be expanded with tool execution, human-in-the-loop, etc.
    workflow.add_edge("draft_strategist", END)
    return workflow.compile()


class DraftingService:
    _instance = None
    # _initialized = False # Initialization status tracked by instance attributes

    def __new__(cls, *args, **kwargs):  # Make generic
        if not cls._instance:
            cls._instance = super(DraftingService, cls).__new__(cls)
            cls._instance._initialized_service = (
                False  # Use a distinct attr to track full init
            )
        return cls._instance

    def __init__(self, settings_param: Optional[Settings] = None):
        if hasattr(self, "_initialized_service") and self._initialized_service:
            return

        current_settings = (
            settings_param if settings_param is not None else get_settings()
        )

        self.settings: Settings = current_settings
        self.output_dir = DEFAULT_DRAFT_OUTPUT_DIR
        logger.info(
            f"DraftingService initializing LLM with API base: {self.settings.OLLAMA_BASE_URL}"
        )
        self.llm_client = ChatOllama(
            model=self.settings.OLLAMA_MODEL_NAME,
            temperature=self.settings.LLM_TEMPERATURE,
            base_url=self.settings.OLLAMA_BASE_URL,
            format="json",
            # Add other parameters like mirostat if needed for better structured output
        )
        logger.info("DraftingService LLM client initialized successfully.")

        # Initialize DocumentLoaderService here, passing settings
        # Assuming DocumentLoaderService also takes settings or uses get_settings()
        try:
            # self.document_loader = get_document_loader_service() # This might create circular dependency if DocumentLoader also uses DraftingService
            # For now, let draft_agent_node handle doc loading if it needs it through a separate mechanism or service.
            logger.info(
                "DraftingService: DocumentLoaderService initialized."
            )  # Placeholder
        except ImportError as e:
            logger.error(
                f"DraftingService: Failed to import DocumentLoaderService: {e}"
            )
            # Handle appropriately, maybe raise or set self.document_loader to None

        # Compile the graph
        self.graph = create_draft_graph()
        logger.info("DraftingService: Draft agent graph compiled.")
        self._initialized_service = True

    def _get_content_from_doc_id(self, document_id: str) -> Optional[str]:
        """Helper to fetch content for a given document_id using DocumentLoaderService."""
        try:
            # Use the injected document loader service
            # The method in DocumentLoaderService is load_document_content_by_id
            content = self.document_loader.load_document_content_by_id(document_id)
            if content:
                return content
            logger.warning(
                f"No content found for document_id: {document_id} by DraftingService's loader"
            )
            return None
        except Exception as e:
            logger.error(
                f"DraftingService: Error loading content for document_id {document_id}: {e}",
                exc_info=True,
            )
            return None

    def _build_llm_context(self, request: DraftStrategyNoteRequest) -> str:
        contexts = []
        if request.claim_summary:
            contexts.append(f"Claim Summary:\n{request.claim_summary}")

        if request.document_ids:
            doc_contents = []
            for doc_id in request.document_ids:
                content = self._get_content_from_doc_id(doc_id)
                if content:
                    doc_contents.append(
                        f"Content from Document ID '{doc_id}':\n{content}"
                    )
            if doc_contents:
                contexts.append("\n\n---\n\n".join(doc_contents))

        if request.qa_history:
            qa_strings = []
            for i, pair in enumerate(request.qa_history):
                qa_strings.append(f"Q&A {i+1}:\nQ: {pair.question}\nA: {pair.answer}")
            if qa_strings:
                contexts.append("Relevant Q&A History:\n" + "\n\n".join(qa_strings))

        if request.additional_criteria:
            contexts.append(
                f"Additional Instructions/Criteria for Drafting:\n{request.additional_criteria}"
            )

        if not contexts:
            # This should ideally be caught by Pydantic validator, but as a fallback.
            logger.warning("No context provided for drafting strategy note.")
            # Raise an error that the endpoint can catch and turn into a 400
            raise ValueError(
                "Insufficient context to draft a strategy note. Please provide summary, documents, Q&A, or criteria."
            )

        return "\n\n---\n\n".join(contexts)

    def generate_strategy_note_text(self, context: str) -> str:
        # This prompt needs to be carefully engineered for good results.
        template = """
        You are an AI assistant tasked with drafting a Claim Strategy Note based on the provided information.
        The note should be comprehensive, well-structured, and actionable.
        Consider including sections such as: Introduction/Background, Key Findings, Strengths, Weaknesses, Potential Risks, Recommended Strategy, Next Steps.
        Adapt the structure and content based on the specifics of the provided context.

        Provided Context:
        ---BEGIN CONTEXT---
        {context_for_llm}
        ---END CONTEXT---

        Draft of Claim Strategy Note:
        """
        prompt = ChatPromptTemplate.from_template(template)
        chain = prompt | self.llm_client | StrOutputParser()

        try:
            logger.info(
                f"Requesting strategy note draft from LLM. Context length: {len(context)}"
            )
            draft_text = chain.invoke({"context_for_llm": context})
            logger.info("Successfully generated strategy note draft text from LLM.")
            return draft_text
        except Exception as e:
            logger.error(
                f"LLM call failed during strategy note drafting: {e}", exc_info=True
            )
            raise HTTPException(
                status_code=500,
                detail="Failed to generate strategy note text due to LLM error.",
            )

    def create_docx_from_text(
        self, *, text: str, filename_suggestion: str | None = None
    ) -> Path:
        """
        Creates a .docx file from the given text and saves it to data/outputs.
        Generates a safe filename, avoiding problematic characters and ensuring uniqueness.
        """
        # 1️⃣ NORMALISE FIRST, *THEN* SPLIT
        cleaned_suggestion = os.path.basename(
            filename_suggestion or "strategy_note.docx"
        )
        original_base_name, ext = os.path.splitext(cleaned_suggestion)

        # Ensure extension is .docx, default if not provided or incorrect
        if not ext.lower() == ".docx":
            ext = ".docx"
            # If the original_base_name was just an extension (e.g. ".txt"),
            # and we are forcing .docx, we should clear original_base_name
            # or we might end up with "txt.docx" if original_base_name was "txt"
            if original_base_name.lower() + ext.lower() == cleaned_suggestion.lower():
                original_base_name = ""

        # Sanitize the base name
        # Remove non-alphanumeric characters (except underscore and hyphen)
        base_name_candidate = re.sub(r"[^\w\-]+", "_", original_base_name)
        # Replace multiple underscores/hyphens with a single one
        base_name_candidate = re.sub(r"[_\-]+", "_", base_name_candidate)
        base_name_candidate = base_name_candidate.strip("_")

        # 2️⃣ FALL BACK IF THE SANITISED NAME IS BLANK **OR JUST THE EXTENSION**
        if (
            not base_name_candidate.strip("_")
            or base_name_candidate.lower() == ext.lstrip(".").lower()
        ):
            final_safe_base_name = f"strategy_note_{uuid.uuid4().hex[:8]}"
        else:
            final_safe_base_name = base_name_candidate

        doc = DocxDocument()
        doc.add_heading("Claim Strategy Note", level=1)

        for para_text in text.split("\n\n"):
            if para_text.strip():
                doc.add_paragraph(para_text.strip())

        safe_filename = final_safe_base_name + ext

        # Limit total length
        max_len = 250  # A reasonable max length for a filename
        if len(safe_filename) > max_len:
            base_max_len = max_len - len(ext)
            final_safe_base_name = final_safe_base_name[:base_max_len]
            safe_filename = final_safe_base_name + ext
            if not final_safe_base_name:  # Re-check if it became empty after truncation
                safe_filename = f"strategy_note_{uuid.uuid4().hex[:8]}.docx"

        output_path = os.path.join(self.output_dir, safe_filename)

        try:
            doc.save(output_path)
            logger.info(f"Strategy note DOCX saved to: {output_path}")
            return Path(output_path)
        except Exception as e:
            logger.error(
                f"Failed to save DOCX file to {output_path}: {e}", exc_info=True
            )
            raise HTTPException(
                status_code=500,
                detail="Failed to save the generated DOCX strategy note.",
            )


# Factory for FastAPI and direct calls
def get_drafting_service() -> DraftingService:
    """Returns a singleton instance of the DraftingService."""
    # __init__ will handle fetching settings if called for the first time.
    return DraftingService()


# Dependency for FastAPI if you want to explicitly show it uses Depends (optional)
# def get_drafting_service_dependency(
#     settings: Settings = Depends(get_settings)
# ) -> DraftingService:
#     return DraftingService(settings_param=settings) # Pass resolved settings
